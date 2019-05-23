import os
import fitz
import time
import requests
import re
from django.conf import settings
from . import translate_func
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn

# 正则匹配参考文献
def is_reference(target):
    return re.match(r'references', target, re.I)

# 正则匹配图片标注
def is_figure(target):
    return re.match(r'fig\..\.', target, re.I)

# 翻译文献到新的pdf以及word中
def trans_pdf(file_name, path):
    t0 = time.time()
    cur_pdf = fitz.open(path)  # 待翻译的pdf
    new_pdf = fitz.open()  # 翻译完成后要写入的pdf
    new_docx = Document()  # 翻译完成后要写入的docx
    new_docx.styles['Normal'].font.name = u'宋体'  # 设置翻译完成后的字体
    new_docx.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')  # 设置翻译完成后的字体
    i = 0  # 定义页面数的递增
    bytes_array = 0
    for cur_page in cur_pdf:
        img_list = cur_page.getImageList()  # 获取当前页面的图片对象
        imgcount = 0
        for img in img_list:  # 获取当前页面的图像列表
            pix_temp = fitz.Pixmap(cur_pdf, img[0])
            print('当前页面的图像', pix_temp)
            imgcount += 1
            new_name = "图片{}.png".format(imgcount)  # 生成图片的名称
            pix_temp.writeImage(os.path.join(settings.BASE_DIR, 'trans', 'output_file', new_name))
            # bytes_array = pix_temp.getImageData('png')#可以不输出图片再写入新的pdf，通过byte
            # print(pix_temp.getImageData('png'))
            pix_temp = None  # 释放资源
        blks = cur_page.getTextBlocks(images=True)  # read text blocks of input page
        new_page = new_pdf.newPage(-1, width=cur_page.MediaBoxSize[0],
                                height=cur_page.MediaBoxSize[1])  # 创建一个新的页面与之前的页面相同大小
        img = new_page.newShape()  # prepare /Contents object
        disp = fitz.Rect(cur_page.CropBoxPosition, cur_page.CropBoxPosition)
        print('disp:', disp)
        croprect = cur_page.rect + disp
        # img.drawRect(croprect)#画出整个页面的矩形
        # img.finish(color=gray, fill=gray)#填充颜色
        begin = (0, 0, 0, 0)  # 记录初始值
        end = (0, 0, 0, 0)  # 记录终结值
        flag = 0  # 记录当前的循
        reference_flag = 0  # 判断是否在参考文献之后
        blks.append((1, 2, 3, 6))
        content = ""
        imgcount = 0
        fonts = 9
        for num in range(len(blks)):  # loop through the blocks
            # 如果是本页面最后一个块,直接结束,因为最后一个是方便计算自己添加的。
            if num == len(blks) - 1:
                break
            # 如果这个块里放的是图像.
            if blks[num][-1] == 1:
                imgcount += 1
                print('blks::', blks[num])
                img_r = blks[num][:4]  # 图片要放置位置的坐标
                path = os.path.join(settings.BASE_DIR, 'trans', 'output_file',
                                    '图片{}.png'.format(imgcount))  # 当前页面第几个图片的位置
                img = open(path, "rb").read()  # 输入流
                new_page.insertImage(img_r, stream=img, keep_proportion=True)  # 输入到新的pdf页面对应位置
                new_docx.add_picture(path, width=Inches(2))  # 设置图片保存的宽度
                try:
                    pass
                    os.remove(path)  # 输入到新的pdf之后就移除
                except:
                    print('删除失败！！')
                continue # 跳过下面的插入翻译后文字的过程

            # 设置默认字体大小以及位置
            if i == 0:  # 当前是第一页的话
                if num == 0 or num == 1:
                    fonts = 15
                    text_pos = fitz.TEXT_ALIGN_CENTER  # 一般论文前面的标题,作者,机构名等要居中
                elif num == 2:
                    fonts = 10
                    text_pos = fitz.TEXT_ALIGN_CENTER  # 一般论文前面的标题,作者,机构名等要居中
                elif num == 3:
                    fonts = 10
                    text_pos = fitz.TEXT_ALIGN_CENTER  # 一般论文前面的标题,作者,机构名等要居中
                else:
                    fonts = 10
                    text_pos = fitz.TEXT_ALIGN_LEFT  # 设置文字在当前矩阵中的位置靠左排列
            else:
                fonts = 10
                text_pos = fitz.TEXT_ALIGN_LEFT  # 设置文字在当前矩阵中的位置靠左排列
            # 目的为了记录起始块坐标
            if num == 0:
                begin = blks[0][:4]
                content = blks[0][4].replace("\n", " ")
            # 矩形块，b[0]b[1]为左上角的坐标，b[2]b[3]为右下角的坐标
            r = fitz.Rect(blks[num][:4])
            # 如果不是倒数第一个块，则进入此循环
            if num < len(blks) - 1:
                # 两个块y轴距离很近的话，这里以0.6为界，这里判断当前数的右下角的坐标y值
                if (abs(blks[num + 1][1] - blks[num][3]) <= 0.6 and abs(
                        blks[num + 1][1] - blks[num][3]) >= 0):
                    # 当前块在参考文献之后
                    if reference_flag == 1:
                        trans_pragraph = blks[num][4].replace("\n", " ")
                        res = translate_func.google_translate(trans_pragraph).replace(' ', '')
                        new_page.insertTextbox(r, res, fontname="song", fontfile=os.path.join(settings.BASE_DIR,
                                                                                              'trans/static/fonts/SimSun.ttf'),
                                               fontsize=7, align=text_pos)  #
                    # 其它情况
                    else:
                        flag = 1  #
                        # 记录最后的矩形坐标，目的为了取出最后的右下角坐标点
                        end = blks[num + 1][:4]
                        content += blks[num + 1][4].replace("\n", " ")
                        # print('content::',content)

                # 两个块y轴距离远的的时候
                else:
                    if flag == 1:
                        # img.drawRect(fitz.Rect(end[0],begin[1],end[2],end[3]))
                        res = translate_func.google_translate(content).replace(' ', '')  # 翻译结果去掉汉字中的空格
                        new_docx.add_paragraph(res)  # 添加到新的docx文档中
                        # print('content:',content)
                        # print(res)
                        # fitz.Rect(end[0],begin[1],end[2],end[3])为新扩展的矩形框坐标
                        if begin[2] > end[2]:  # 如果起始点的右下角x坐标小于结束点的右下角x坐标
                            new_page.insertTextbox(fitz.Rect(end[0], begin[1], begin[2], end[3]), res, fontname="song",
                                                fontfile=os.path.join(settings.BASE_DIR,
                                                                      'trans/static/fonts/SimSun.ttf'),
                                                fontsize=fonts, align=text_pos)
                        else:
                            new_page.insertTextbox(fitz.Rect(end[0], begin[1], end[2], end[3]), res, fontname="song",
                                                fontfile=os.path.join(settings.BASE_DIR,
                                                                      'trans/static/fonts/SimSun.ttf'),
                                                fontsize=fonts, align=text_pos)
                        flag = 0
                    else:
                        # img.drawRect(r)
                        trans_pragraph = blks[num][4].replace("\n", " ")  # 将待翻译的句子换行换成空格
                        if is_figure(trans_pragraph.replace(' ','')):  # 将该块的判断是否是图片标注
                            res = translate_func.google_translate(trans_pragraph).replace(' ', '')  # 翻译结果去掉汉字中的空格
                            new_page.insertTextbox(r, res, fontname="song", fontfile=os.path.join(settings.BASE_DIR,
                                                                                               'trans/static/fonts/SimSun.ttf'),
                                                fontsize=7, align=fitz.TEXT_ALIGN_CENTER)
                        # 标记在这里之后的都是参考文献
                        elif is_reference(trans_pragraph.replace(' ','')):
                            reference_flag = 1
                            new_page.insertTextbox(r, '参考文献', fontname="song", fontfile=os.path.join(settings.BASE_DIR,
                                                                                               'trans/static/fonts/SimSun.ttf'),
                                                fontsize=fonts, align=text_pos)
                        else:
                            # 翻译结果去掉汉字中的空格
                            res = translate_func.google_translate(trans_pragraph).replace(' ', '')
                            # 添加到新的docx文档中
                            new_docx.add_paragraph(res)
                            if reference_flag == 1:
                                new_page.insertTextbox(r, res, fontname="song", fontfile=os.path.join(settings.BASE_DIR,
                                                                                                      'trans/static/fonts/SimSun.ttf'),
                                                       fontsize=7, align=text_pos)  #
                            else:

                                new_page.insertTextbox(r, res, fontname="song", fontfile=os.path.join(settings.BASE_DIR,
                                                                                               'trans/static/fonts/SimSun.ttf'),
                                                fontsize=fonts, align=text_pos)  #
                    # 记录起始矩形坐标
                    begin = blks[num + 1][:4]
                    try:
                        content = blks[num + 1][4].replace("\n", " ")
                        # print('content:::',content)
                    except:
                        print('记录content失败！')
                    # img.finish(width=0.3)
                    # img.commit()
        i += 1
    # 文件保存
    new_file_name = os.path.join(settings.BASE_DIR, 'trans', 'output_file', 'translated_' + file_name)  # 翻译后的pdf保存路径
    new_docx_name = os.path.join(settings.BASE_DIR, 'trans', 'output_file',
                                 'translated_' + file_name[:-4] + '.docx')  # 翻译后的docx保存路径
    new_docx.save(new_docx_name)  # 保存翻译后的docx
    new_pdf.save(new_file_name, garbage=4, deflate=True, clean=True)  # 保存翻译后的pdf
    t1 = time.time()
    print("Total translation time: %g sec" % (t1 - t0))
